from pathlib import Path
from typing import List, Dict, Any
import logging
from datetime import datetime

try:
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDoc
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    DocxDocument = None
    DocxDoc = None
    Table = None
    Paragraph = None

from core.base_extractor import ExtractorPlugin
from core.data_models import Document, ProcessingInfo

logger = logging.getLogger(__name__)


class DocxExtractor(ExtractorPlugin):
    """DOCX document extractor for paragraphs and tables."""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    def can_handle(self, file_path: Path) -> bool:
        if not file_path.suffix.lower() in ['.docx', '.doc']:
            return False
        
        if not HAS_DOCX:
            logger.warning(f"Cannot handle DOCX {file_path}: Missing python-docx dependency")
            return False
            
        return True
    
    def extract(self, file_path: Path) -> Document:
        """Extract text and metadata from DOCX file."""
        if not HAS_DOCX:
            raise RuntimeError("DOCX processing requires python-docx")
            
        doc = Document.from_file(file_path)
        
        processing_info = ProcessingInfo(
            extraction_method='python-docx',
            chunking_strategy='none',
            processing_timestamp=datetime.now()
        )
        
        try:
            docx_doc = DocxDocument(str(file_path))
            
            # Extract content with structure
            content, structure = self._extract_content_with_structure(docx_doc)
            doc.content = content
            doc.structure = structure
            
            # Extract metadata
            metadata_updates = self._extract_metadata(docx_doc)
            for key, value in metadata_updates.items():
                setattr(doc.metadata, key, value)
            
            # Update character and word counts
            doc.metadata.char_count = len(content)
            doc.metadata.word_count = len(content.split()) if content else 0
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX content from {file_path}: {e}")
            processing_info.errors.append(str(e))
            doc.content = ""
            doc.structure = {}
        
        return doc
    
    def _extract_content_with_structure(self, docx_doc: DocxDoc) -> tuple[str, Dict[str, Any]]:
        """Extract content while preserving document structure."""
        content_parts = []
        structure = {
            'paragraphs': [],
            'tables': [],
            'headers': [],
            'sections': []
        }
        
        paragraph_count = 0
        table_count = 0
        
        for element in docx_doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                paragraph = None
                for p in docx_doc.paragraphs:
                    if p._element is element:
                        paragraph = p
                        break
                
                if paragraph:
                    text = paragraph.text.strip()
                    if text:
                        # Check if it's a heading
                        style_name = paragraph.style.name if paragraph.style else ""
                        is_heading = 'Heading' in style_name or 'Title' in style_name
                        
                        if is_heading:
                            content_parts.append(f"\n## {text}\n")
                            structure['headers'].append({
                                'text': text,
                                'style': style_name,
                                'level': self._get_heading_level(style_name),
                                'paragraph_index': paragraph_count
                            })
                        else:
                            content_parts.append(text)
                        
                        structure['paragraphs'].append({
                            'index': paragraph_count,
                            'text': text,
                            'style': style_name,
                            'is_heading': is_heading,
                            'char_count': len(text),
                            'word_count': len(text.split())
                        })
                        
                        paragraph_count += 1
                        
            elif element.tag.endswith('tbl'):  # Table
                table = None
                for t in docx_doc.tables:
                    if t._element is element:
                        table = t
                        break
                
                if table:
                    table_text = self._extract_table_text(table)
                    content_parts.append(f"\n[Table {table_count + 1}]\n{table_text}\n")
                    
                    structure['tables'].append({
                        'index': table_count,
                        'rows': len(table.rows),
                        'cols': len(table.columns) if table.rows else 0,
                        'text': table_text
                    })
                    
                    table_count += 1
        
        return '\n\n'.join(content_parts), structure
    
    def _extract_table_text(self, table: Table) -> str:
        """Extract text from a DOCX table."""
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            table_data.append(' | '.join(row_data))
        
        return '\n'.join(table_data)
    
    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if 'Title' in style_name:
            return 1
        elif 'Heading 1' in style_name:
            return 1
        elif 'Heading 2' in style_name:
            return 2
        elif 'Heading 3' in style_name:
            return 3
        elif 'Heading 4' in style_name:
            return 4
        elif 'Heading 5' in style_name:
            return 5
        elif 'Heading 6' in style_name:
            return 6
        else:
            # Try to extract number from style name
            import re
            match = re.search(r'Heading (\d+)', style_name)
            if match:
                return int(match.group(1))
            return 1
    
    def _extract_metadata(self, docx_doc: DocxDoc) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata = {}
        
        try:
            core_props = docx_doc.core_properties
            
            if core_props.title:
                metadata['title'] = core_props.title
            
            if core_props.author:
                metadata['author'] = core_props.author
            
            if core_props.subject:
                metadata['subject'] = core_props.subject
            
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords.split(',') if ',' in core_props.keywords else [core_props.keywords]
            
            if core_props.comments:
                metadata['comments'] = core_props.comments
            
            if core_props.category:
                metadata['category'] = core_props.category
            
            if core_props.created:
                metadata['created_date'] = core_props.created
            
            if core_props.modified:
                metadata['modified_date'] = core_props.modified
            
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            
            if core_props.revision:
                metadata['revision'] = core_props.revision
                
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
        
        return metadata